"""
Python helper service for n8n workflows.

Handles tasks that are difficult/impossible in n8n's restricted Code node sandbox:
- .docx and .pdf text extraction
- Safe JSON body construction for Azure OpenAI chat completions
- HTML to PDF conversion for proposal generation
"""

import base64
import io
import json
import logging
import re
import urllib.request
import urllib.error
from typing import Optional

from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel

app = FastAPI(title="n8n Python Helper", version="1.0.0")
logger = logging.getLogger("python-helper")
logging.basicConfig(level=logging.INFO)


# ── Health ──────────────────────────────────────────────────────────────────


@app.get("/health")
def health():
    return {"status": "ok"}


# ── Text extraction ─────────────────────────────────────────────────────────


class ExtractTextRequest(BaseModel):
    """Accepts base64-encoded file content + filename to determine type."""
    base64Content: str
    fileName: str
    mimeType: Optional[str] = None


class ExtractTextResponse(BaseModel):
    text: str
    charCount: int
    method: str  # "docx" | "pdf" | "utf8-fallback"


@app.post("/extract-text", response_model=ExtractTextResponse)
def extract_text(req: ExtractTextRequest):
    """Extract plain text from .docx, .pdf, or fall back to UTF-8 decode."""
    try:
        raw = base64.b64decode(req.base64Content)
    except Exception as e:
        raise HTTPException(400, f"Invalid base64: {e}")

    name_lower = req.fileName.lower()
    mime = (req.mimeType or "").lower()

    # ── .docx ───────────────────────────────────────────────────────────
    if name_lower.endswith(".docx") or "wordprocessingml" in mime:
        try:
            from docx import Document

            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append("\t".join(cells))
            text = "\n".join(paragraphs)
            return ExtractTextResponse(text=text, charCount=len(text), method="docx")
        except Exception as e:
            logger.warning("docx extraction failed, trying fallback: %s", e)

    # ── .pdf ────────────────────────────────────────────────────────────
    if name_lower.endswith(".pdf") or "pdf" in mime:
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            pages = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            text = "\n\n".join(pages)
            return ExtractTextResponse(text=text, charCount=len(text), method="pdf")
        except Exception as e:
            logger.warning("pdf extraction failed, trying fallback: %s", e)

    # ── Fallback: treat as UTF-8 text ───────────────────────────────────
    try:
        text = raw.decode("utf-8", errors="replace")
        # Strip any null bytes or control chars except newline/tab
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        return ExtractTextResponse(text=text, charCount=len(text), method="utf8-fallback")
    except Exception as e:
        raise HTTPException(500, f"All extraction methods failed: {e}")


@app.post("/extract-text-file", response_model=ExtractTextResponse)
async def extract_text_file(file: UploadFile = File(...)):
    """Extract plain text from an uploaded file (.docx, .pdf, or plain text)."""
    raw = await file.read()
    name_lower = (file.filename or "unknown").lower()
    mime = (file.content_type or "").lower()

    # ── .docx ───────────────────────────────────────────────────────────
    if name_lower.endswith(".docx") or "wordprocessingml" in mime:
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        paragraphs.append("\t".join(cells))
            text = "\n".join(paragraphs)
            return ExtractTextResponse(text=text, charCount=len(text), method="docx")
        except Exception as e:
            logger.warning("docx extraction failed, trying fallback: %s", e)

    # ── .pdf ────────────────────────────────────────────────────────────
    if name_lower.endswith(".pdf") or "pdf" in mime:
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            pages = [p.extract_text() for p in reader.pages if p.extract_text()]
            text = "\n\n".join(pages)
            return ExtractTextResponse(text=text, charCount=len(text), method="pdf")
        except Exception as e:
            logger.warning("pdf extraction failed, trying fallback: %s", e)

    # ── Fallback: treat as UTF-8 text ───────────────────────────────────
    try:
        text = raw.decode("utf-8", errors="replace")
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        return ExtractTextResponse(text=text, charCount=len(text), method="utf8-fallback")
    except Exception as e:
        raise HTTPException(500, f"All extraction methods failed: {e}")


# ── Chat body builder ───────────────────────────────────────────────────────


class ChatMessage(BaseModel):
    role: str  # "system" | "user" | "assistant"
    content: str


class BuildChatBodyRequest(BaseModel):
    """Build a properly JSON-escaped Azure OpenAI chat completions body."""
    messages: list[ChatMessage]
    temperature: float = 0.1
    max_tokens: int = 2000
    response_format: Optional[dict] = None


@app.post("/build-chat-body")
def build_chat_body(req: BuildChatBodyRequest):
    """
    Returns a ready-to-send JSON object for Azure OpenAI chat completions.
    All string values are properly escaped via json.dumps().
    """
    body = {
        "messages": [{"role": m.role, "content": m.content} for m in req.messages],
        "temperature": req.temperature,
        "max_tokens": req.max_tokens,
    }
    if req.response_format:
        body["response_format"] = req.response_format
    # Return the body directly — FastAPI's JSON serializer handles escaping
    return body


# ── Excel summary extraction ────────────────────────────────────────────────


class ExcelSummaryResponse(BaseModel):
    summary: str
    worksheetCount: int
    totalRows: int


@app.post("/extract-excel-summary", response_model=ExcelSummaryResponse)
async def extract_excel_summary(file: UploadFile = File(...)):
    """Extract worksheet names and first 20 rows from an uploaded .xlsx file."""
    raw = await file.read()
    try:
        from openpyxl import load_workbook

        wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        parts = []
        total_rows = 0

        for ws in wb.worksheets:
            rows = list(ws.iter_rows(max_row=21, values_only=True))  # header + 20
            row_count = ws.max_row or len(rows)
            col_count = ws.max_column or 0
            total_rows += row_count

            section = f'Worksheet: "{ws.title}" ({row_count} rows, {col_count} cols)'

            if rows:
                # First row as headers
                headers = [str(c) if c is not None else "" for c in rows[0]]
                section += f"\nHeaders: {', '.join(headers)}"

                # Data rows (up to 20)
                for i, row in enumerate(rows[1:21], start=1):
                    cells = [str(c) if c is not None else "" for c in row]
                    section += f"\nRow {i}: {', '.join(cells)}"

            parts.append(section)

        wb.close()
        summary = "\n\n".join(parts)
        return ExcelSummaryResponse(
            summary=summary,
            worksheetCount=len(wb.sheetnames),
            totalRows=total_rows,
        )
    except Exception as e:
        logger.error("Excel summary extraction failed: %s", e)
        raise HTTPException(500, f"Excel extraction failed: {e}")


# ── Batch attachment extraction ─────────────────────────────────────────────


class AttachmentItem(BaseModel):
    name: str
    contentType: Optional[str] = None
    contentBytes: str  # base64-encoded


class ExtractAttachmentsRequest(BaseModel):
    attachments: list[AttachmentItem]


class ExtractAttachmentsResponse(BaseModel):
    summary: str
    attachmentCount: int
    analyzedCount: int


ANALYZABLE_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xls", ".txt", ".csv"}


@app.post("/extract-attachments-summary", response_model=ExtractAttachmentsResponse)
def extract_attachments_summary(req: ExtractAttachmentsRequest):
    """
    Extract text from multiple email attachments in one call.
    Accepts base64-encoded file content for each attachment.
    Supports .docx, .pdf, .xlsx, .txt, .csv files.
    """
    parts = []
    analyzed = 0

    for att in req.attachments:
        name_lower = att.name.lower()
        ext = "." + name_lower.rsplit(".", 1)[-1] if "." in name_lower else ""

        if ext not in ANALYZABLE_EXTENSIONS:
            parts.append(f"[{att.name}]: (skipped — unsupported file type)")
            continue

        try:
            raw = base64.b64decode(att.contentBytes)
        except Exception as e:
            parts.append(f"[{att.name}]: (error decoding base64: {e})")
            continue

        mime = (att.contentType or "").lower()

        # ── .docx ───────────────────────────────────────────────────
        if ext == ".docx" or "wordprocessingml" in mime:
            try:
                from docx import Document
                doc = Document(io.BytesIO(raw))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if cells:
                            paragraphs.append("\t".join(cells))
                text = "\n".join(paragraphs)
                # Truncate to 3000 chars per attachment
                if len(text) > 3000:
                    text = text[:3000] + "... (truncated)"
                parts.append(f"[{att.name}]:\n{text}")
                analyzed += 1
                continue
            except Exception as e:
                logger.warning("docx extraction failed for %s: %s", att.name, e)

        # ── .pdf ────────────────────────────────────────────────────
        if ext == ".pdf" or "pdf" in mime:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(raw))
                pages = [p.extract_text() for p in reader.pages if p.extract_text()]
                text = "\n\n".join(pages)
                if len(text) > 3000:
                    text = text[:3000] + "... (truncated)"
                parts.append(f"[{att.name}]:\n{text}")
                analyzed += 1
                continue
            except Exception as e:
                logger.warning("pdf extraction failed for %s: %s", att.name, e)

        # ── .xlsx ───────────────────────────────────────────────────
        if ext in (".xlsx", ".xls") or "spreadsheet" in mime:
            try:
                from openpyxl import load_workbook
                wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
                ws_parts = []
                for ws in wb.worksheets:
                    rows = list(ws.iter_rows(max_row=21, values_only=True))
                    row_count = ws.max_row or len(rows)
                    col_count = ws.max_column or 0
                    section = f'  Worksheet: "{ws.title}" ({row_count} rows, {col_count} cols)'
                    if rows:
                        headers = [str(c) if c is not None else "" for c in rows[0]]
                        section += f"\n  Headers: {', '.join(headers)}"
                        for i, row in enumerate(rows[1:21], start=1):
                            cells = [str(c) if c is not None else "" for c in row]
                            section += f"\n  Row {i}: {', '.join(cells)}"
                    ws_parts.append(section)
                wb.close()
                text = "\n".join(ws_parts)
                if len(text) > 3000:
                    text = text[:3000] + "... (truncated)"
                parts.append(f"[{att.name}]:\n{text}")
                analyzed += 1
                continue
            except Exception as e:
                logger.warning("xlsx extraction failed for %s: %s", att.name, e)

        # ── .txt / .csv / fallback ──────────────────────────────────
        try:
            text = raw.decode("utf-8", errors="replace")
            text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
            if len(text) > 3000:
                text = text[:3000] + "... (truncated)"
            parts.append(f"[{att.name}]:\n{text}")
            analyzed += 1
        except Exception as e:
            parts.append(f"[{att.name}]: (extraction failed: {e})")

    summary = "\n\n".join(parts) if parts else "(no analyzable attachments)"
    return ExtractAttachmentsResponse(
        summary=summary,
        attachmentCount=len(req.attachments),
        analyzedCount=analyzed,
    )


# ── HTML to PDF conversion ──────────────────────────────────────────────────


class HtmlToPdfRequest(BaseModel):
    """Convert HTML string to PDF bytes (returned as base64 or binary)."""
    html: str
    filename: Optional[str] = "document.pdf"


@app.post("/html-to-pdf")
def html_to_pdf(req: HtmlToPdfRequest):
    """
    Convert an HTML string to a PDF file.
    Returns the PDF as a binary response with Content-Type: application/pdf.
    """
    try:
        from weasyprint import HTML

        pdf_bytes = HTML(string=req.html).write_pdf()
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{req.filename}"',
            },
        )
    except Exception as e:
        logger.error("HTML to PDF conversion failed: %s", e)
        raise HTTPException(500, f"PDF conversion failed: {e}")


@app.post("/html-to-pdf-base64")
def html_to_pdf_base64(req: HtmlToPdfRequest):
    """
    Convert an HTML string to a PDF and return as base64-encoded string.
    Useful for n8n workflows that need to upload the PDF via HTTP to SharePoint.
    """
    try:
        from weasyprint import HTML

        pdf_bytes = HTML(string=req.html).write_pdf()
        b64 = base64.b64encode(pdf_bytes).decode("ascii")
        return {
            "base64Content": b64,
            "filename": req.filename,
            "sizeBytes": len(pdf_bytes),
            "contentType": "application/pdf",
        }
    except Exception as e:
        logger.error("HTML to PDF base64 conversion failed: %s", e)
        raise HTTPException(500, f"PDF conversion failed: {e}")


# ── HTML to DOCX conversion ────────────────────────────────────────────────


class HtmlToDocxRequest(BaseModel):
    """Convert structured proposal JSON to a Word document."""
    proposal: dict
    lead: dict
    filename: Optional[str] = "document.docx"


@app.post("/html-to-docx-base64")
def html_to_docx_base64(req: HtmlToDocxRequest):
    """
    Build a professional Word document from proposal JSON data.
    Returns the .docx as a base64-encoded string for SharePoint upload.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # ── Page margins ────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ── Style colors ────────────────────────────────────────────
        brand_green = RGBColor(0x9A, 0xB2, 0x3B)
        brand_brown = RGBColor(0x63, 0x5E, 0x59)
        text_gray = RGBColor(0x33, 0x33, 0x33)
        light_gray = RGBColor(0x99, 0x99, 0x99)

        proposal = req.proposal
        lead = req.lead

        # ── Header ──────────────────────────────────────────────────
        title = doc.add_heading("Executive Catering CT", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = brand_green
            run.font.size = Pt(28)

        subtitle = doc.add_paragraph("Catering Proposal")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.color.rgb = light_gray
            run.font.size = Pt(14)

        # ── Date and ref ────────────────────────────────────────────
        from datetime import datetime
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
        date_run.font.color.rgb = light_gray
        date_run.font.size = Pt(10)

        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"Ref: {lead.get('leadId', '')}")
        ref_run.font.color.rgb = light_gray
        ref_run.font.size = Pt(10)

        # ── Greeting ────────────────────────────────────────────────
        first = lead.get("firstName", "")
        last = lead.get("lastName", "")
        greeting_name = f"{first} {last}".strip() or "Valued Client"
        greeting = doc.add_heading(f"Dear {greeting_name},", level=2)
        for run in greeting.runs:
            run.font.color.rgb = brand_brown

        # ── Company intro ───────────────────────────────────────────
        if proposal.get("companyIntro"):
            doc.add_paragraph(proposal["companyIntro"])

        # ── Event summary ───────────────────────────────────────────
        h = doc.add_heading("Event Summary", level=2)
        for run in h.runs:
            run.font.color.rgb = brand_brown
        if proposal.get("eventSummary"):
            doc.add_paragraph(proposal["eventSummary"])

        # ── Menu options ────────────────────────────────────────────
        h = doc.add_heading("Menu Options", level=2)
        for run in h.runs:
            run.font.color.rgb = brand_brown

        for i, opt in enumerate(proposal.get("menuOptions", []), 1):
            tier_heading = doc.add_heading(
                f"Option {i}: {opt.get('tierName', '')} — {opt.get('pricePerPerson', '')}/person",
                level=3
            )
            for run in tier_heading.runs:
                run.font.color.rgb = brand_green

            if opt.get("description"):
                desc = doc.add_paragraph(opt["description"])
                for run in desc.runs:
                    run.font.italic = True

            categories = [
                ("Appetizers", opt.get("appetizers", [])),
                ("Entrées", opt.get("entrees", [])),
                ("Sides", opt.get("sides", [])),
                ("Desserts", opt.get("desserts", [])),
                ("Beverages", opt.get("beverages", [])),
            ]
            for cat_name, items in categories:
                if items:
                    cat_para = doc.add_paragraph()
                    cat_run = cat_para.add_run(f"{cat_name}: ")
                    cat_run.font.bold = True
                    cat_run.font.color.rgb = brand_brown
                    cat_para.add_run(", ".join(items))

        # ── Staffing plan ───────────────────────────────────────────
        h = doc.add_heading("Staffing Plan", level=2)
        for run in h.runs:
            run.font.color.rgb = brand_brown
        if proposal.get("staffingPlan"):
            doc.add_paragraph(proposal["staffingPlan"])

        # ── Equipment & rentals ─────────────────────────────────────
        h = doc.add_heading("Equipment & Rentals", level=2)
        for run in h.runs:
            run.font.color.rgb = brand_brown
        if proposal.get("equipmentAndRentals"):
            doc.add_paragraph(proposal["equipmentAndRentals"])

        # ── Event timeline (table) ──────────────────────────────────
        timeline = proposal.get("timeline", [])
        if timeline:
            h = doc.add_heading("Event Timeline", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_brown

            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Light Shading Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Time"
            hdr[1].text = "Activity"
            for run in hdr[0].paragraphs[0].runs:
                run.font.bold = True
            for run in hdr[1].paragraphs[0].runs:
                run.font.bold = True

            for entry in timeline:
                row = table.add_row().cells
                row[0].text = entry.get("time", "")
                row[1].text = entry.get("activity", "")

        # ── Pricing summary (table) ─────────────────────────────────
        pricing = proposal.get("pricingSummary", {})
        if pricing:
            h = doc.add_heading("Pricing Summary", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_brown

            price_table = doc.add_table(rows=0, cols=2)
            price_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            price_table.style = "Light Shading Accent 1"

            price_items = [
                ("Food & Beverage", pricing.get("foodAndBeverage", "TBD")),
                ("Staffing", pricing.get("staffing", "TBD")),
                ("Equipment & Rentals", pricing.get("equipmentRentals", "Included")),
                ("Delivery & Setup", pricing.get("deliveryAndSetup", "TBD")),
            ]
            for label, value in price_items:
                row = price_table.add_row().cells
                row[0].text = label
                row[1].text = value
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Total row (bold + green)
            total_row = price_table.add_row().cells
            total_run = total_row[0].paragraphs[0].add_run("Estimated Total")
            total_run.font.bold = True
            total_run.font.size = Pt(12)
            total_val = total_row[1].paragraphs[0].add_run(
                pricing.get("estimatedTotal", "TBD")
            )
            total_val.font.bold = True
            total_val.font.size = Pt(12)
            total_val.font.color.rgb = brand_green
            total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Deposit row
            dep_row = price_table.add_row().cells
            dep_row[0].text = "Deposit Required"
            dep_row[1].text = pricing.get("depositRequired", "50%")
            dep_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # ── Terms ───────────────────────────────────────────────────
        h = doc.add_heading("Terms & Conditions", level=2)
        for run in h.runs:
            run.font.color.rgb = brand_brown
        if proposal.get("terms"):
            doc.add_paragraph(proposal["terms"])

        # ── Footer ──────────────────────────────────────────────────
        doc.add_paragraph("")  # spacer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = footer.add_run("Executive Catering CT | Connecticut's Premier Catering Service")
        r1.font.color.rgb = light_gray
        r1.font.size = Pt(9)
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = footer2.add_run("This proposal is valid for 30 days from the date of preparation.")
        r2.font.color.rgb = light_gray
        r2.font.size = Pt(9)

        # ── Serialize to bytes ──────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        b64 = base64.b64encode(docx_bytes).decode("ascii")

        return {
            "base64Content": b64,
            "filename": req.filename,
            "sizeBytes": len(docx_bytes),
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    except Exception as e:
        logger.error("HTML to DOCX conversion failed: %s", e)
        raise HTTPException(500, f"DOCX conversion failed: {e}")


# ── OpsBlueprint proposal to DOCX conversion ──────────────────────────────


@app.post("/opsblueprint-to-docx-base64")
def opsblueprint_to_docx_base64(req: HtmlToDocxRequest):
    """
    Build a professional Word document from OpsBlueprint proposal JSON data.
    Expects the OpsBlueprint proposal schema (scopeOfWork, deliverables,
    automationOpportunities, etc.) rather than the catering schema.
    Returns the .docx as a base64-encoded string for SharePoint upload.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()

        # ── Page margins ────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # ── Brand colors ────────────────────────────────────────────
        brand_blue = RGBColor(0x1E, 0x3A, 0x5F)      # Deep navy
        brand_accent = RGBColor(0x2B, 0x7A, 0xD9)     # Bright blue
        text_dark = RGBColor(0x2D, 0x2D, 0x2D)
        light_gray = RGBColor(0x99, 0x99, 0x99)
        success_green = RGBColor(0x27, 0xAE, 0x60)

        proposal = req.proposal
        lead = req.lead

        # ── Header ──────────────────────────────────────────────────
        title = doc.add_heading("OpsBlueprint", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = brand_blue
            run.font.size = Pt(28)

        subtitle = doc.add_paragraph("Workflow Automation Consulting Proposal")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in subtitle.runs:
            run.font.color.rgb = brand_accent
            run.font.size = Pt(14)

        # ── Date and validity ───────────────────────────────────────
        from datetime import datetime
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(
            f"Prepared: {proposal.get('proposalDate', datetime.now().strftime('%B %d, %Y'))}"
        )
        date_run.font.color.rgb = light_gray
        date_run.font.size = Pt(10)

        if proposal.get("validUntil"):
            valid_para = doc.add_paragraph()
            valid_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            valid_run = valid_para.add_run(f"Valid Until: {proposal['validUntil']}")
            valid_run.font.color.rgb = light_gray
            valid_run.font.size = Pt(10)

        ref_para = doc.add_paragraph()
        ref_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_para.add_run(f"Ref: {lead.get('leadId', '')}")
        ref_run.font.color.rgb = light_gray
        ref_run.font.size = Pt(10)

        # ── Client info ─────────────────────────────────────────────
        contact_name = (
            proposal.get("contactName")
            or lead.get("fullName")
            or "Valued Client"
        )
        company_name = proposal.get("companyName") or lead.get("company", "")

        greeting = doc.add_heading(f"Prepared for: {contact_name}", level=2)
        for run in greeting.runs:
            run.font.color.rgb = brand_blue

        if company_name:
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(company_name)
            company_run.font.size = Pt(12)
            company_run.font.color.rgb = text_dark

        # ── Executive Summary ───────────────────────────────────────
        if proposal.get("executiveSummary"):
            h = doc.add_heading("Executive Summary", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue
            doc.add_paragraph(proposal["executiveSummary"])

        # ── Recommended Tier & Investment ───────────────────────────
        tier = proposal.get("recommendedTier", "")
        investment = proposal.get("estimatedInvestment", "")
        if tier or investment:
            h = doc.add_heading("Recommended Service Tier", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if tier:
                tier_para = doc.add_paragraph()
                tier_label = tier_para.add_run("Tier: ")
                tier_label.font.bold = True
                tier_label.font.color.rgb = text_dark
                tier_value = tier_para.add_run(tier)
                tier_value.font.color.rgb = brand_accent
                tier_value.font.size = Pt(14)
                tier_value.font.bold = True

            if investment:
                inv_para = doc.add_paragraph()
                inv_label = inv_para.add_run("Estimated Investment: ")
                inv_label.font.bold = True
                inv_label.font.color.rgb = text_dark
                inv_value = inv_para.add_run(investment)
                inv_value.font.color.rgb = success_green
                inv_value.font.size = Pt(13)
                inv_value.font.bold = True

        # ── Scope of Work ───────────────────────────────────────────
        scope = proposal.get("scopeOfWork", [])
        if scope:
            h = doc.add_heading("Scope of Work", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if isinstance(scope, list):
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Light Shading Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "Phase"
                hdr[1].text = "Description"
                hdr[2].text = "Duration"
                for cell in hdr:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

                for phase in scope:
                    if isinstance(phase, dict):
                        row = table.add_row().cells
                        row[0].text = phase.get("phase", "")
                        row[1].text = phase.get("description", "")
                        row[2].text = phase.get("duration", "")
                    elif isinstance(phase, str):
                        row = table.add_row().cells
                        row[0].text = phase
                        row[1].text = ""
                        row[2].text = ""
            elif isinstance(scope, str):
                doc.add_paragraph(scope)

        # ── Deliverables ────────────────────────────────────────────
        deliverables = proposal.get("deliverables", [])
        if deliverables:
            h = doc.add_heading("Deliverables", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if isinstance(deliverables, list):
                for item in deliverables:
                    if isinstance(item, str):
                        bullet = doc.add_paragraph(item, style="List Bullet")
                        for run in bullet.runs:
                            run.font.color.rgb = text_dark
                    elif isinstance(item, dict):
                        bullet = doc.add_paragraph(
                            item.get("name", item.get("deliverable", str(item))),
                            style="List Bullet",
                        )
                        for run in bullet.runs:
                            run.font.color.rgb = text_dark
            elif isinstance(deliverables, str):
                doc.add_paragraph(deliverables)

        # ── Automation Opportunities ────────────────────────────────
        opportunities = proposal.get("automationOpportunities", [])
        if opportunities:
            h = doc.add_heading("Automation Opportunities", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if isinstance(opportunities, list):
                table = doc.add_table(rows=1, cols=3)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.style = "Light Shading Accent 1"
                hdr = table.rows[0].cells
                hdr[0].text = "Process"
                hdr[1].text = "Proposed Solution"
                hdr[2].text = "Estimated ROI"
                for cell in hdr:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True

                for opp in opportunities:
                    if isinstance(opp, dict):
                        row = table.add_row().cells
                        row[0].text = opp.get("process", "")
                        row[1].text = opp.get("solution", "")
                        row[2].text = opp.get("estimatedROI", "")

        # ── Timeline ────────────────────────────────────────────────
        timeline = proposal.get("timeline", "")
        if timeline:
            h = doc.add_heading("Project Timeline", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if isinstance(timeline, str):
                doc.add_paragraph(timeline)
            elif isinstance(timeline, list):
                for entry in timeline:
                    if isinstance(entry, dict):
                        item_para = doc.add_paragraph(style="List Bullet")
                        phase_run = item_para.add_run(
                            f"{entry.get('phase', entry.get('time', ''))}: "
                        )
                        phase_run.font.bold = True
                        item_para.add_run(
                            entry.get("description", entry.get("activity", ""))
                        )
                    elif isinstance(entry, str):
                        doc.add_paragraph(entry, style="List Bullet")

        # ── Payment Terms ───────────────────────────────────────────
        payment = proposal.get("paymentTerms", "")
        if payment:
            h = doc.add_heading("Payment Terms", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue
            doc.add_paragraph(payment)

        # ── Next Steps ──────────────────────────────────────────────
        next_steps = proposal.get("nextSteps", [])
        if next_steps:
            h = doc.add_heading("Next Steps", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue

            if isinstance(next_steps, list):
                for i, step in enumerate(next_steps, 1):
                    step_text = step if isinstance(step, str) else str(step)
                    step_para = doc.add_paragraph()
                    num_run = step_para.add_run(f"{i}. ")
                    num_run.font.bold = True
                    num_run.font.color.rgb = brand_accent
                    step_para.add_run(step_text)
            elif isinstance(next_steps, str):
                doc.add_paragraph(next_steps)

        # ── Terms & Conditions ──────────────────────────────────────
        if proposal.get("terms"):
            h = doc.add_heading("Terms & Conditions", level=2)
            for run in h.runs:
                run.font.color.rgb = brand_blue
            doc.add_paragraph(proposal["terms"])

        # ── Footer ──────────────────────────────────────────────────
        doc.add_paragraph("")  # spacer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = footer.add_run(
            "OpsBlueprint | Workflow Automation Consulting"
        )
        r1.font.color.rgb = light_gray
        r1.font.size = Pt(9)
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = footer2.add_run(
            "This proposal is valid for 30 days from the date of preparation."
        )
        r2.font.color.rgb = light_gray
        r2.font.size = Pt(9)

        # ── Serialize to bytes ──────────────────────────────────────
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()
        b64 = base64.b64encode(docx_bytes).decode("ascii")

        return {
            "base64Content": b64,
            "filename": req.filename,
            "sizeBytes": len(docx_bytes),
            "contentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    except Exception as e:
        logger.error("OpsBlueprint DOCX conversion failed: %s", e)
        raise HTTPException(500, f"DOCX conversion failed: {e}")


# ── Generate DOCX and upload to SharePoint in one call ──────────────────────


class DocxAndUploadRequest(BaseModel):
    """Generate a DOCX from proposal data and upload it directly to SharePoint."""
    proposal: dict
    lead: dict
    filename: Optional[str] = "document.docx"
    accessToken: str
    driveId: str
    folderId: str


@app.post("/generate-docx-and-upload")
def generate_docx_and_upload(req: DocxAndUploadRequest):
    """
    Generate a professional Word document from proposal JSON, then upload
    it directly to SharePoint via MS Graph PUT.

    This bypasses n8n's Code node sandbox which cannot send raw binary
    bodies (Buffer objects get serialized to JSON instead of raw bytes).
    """
    # Step 1: Generate the DOCX — detect proposal type and use the right builder
    inner_req = HtmlToDocxRequest(
        proposal=req.proposal,
        lead=req.lead,
        filename=req.filename,
    )

    # Detect OpsBlueprint proposals by their unique fields
    is_opsblueprint = (
        "scopeOfWork" in req.proposal
        or "automationOpportunities" in req.proposal
        or "recommendedTier" in req.proposal
    )

    if is_opsblueprint:
        docx_result = opsblueprint_to_docx_base64(inner_req)
    else:
        docx_result = html_to_docx_base64(inner_req)

    docx_bytes = base64.b64decode(docx_result["base64Content"])
    docx_size = len(docx_bytes)

    logger.info(
        "Generated DOCX: %s (%d bytes), uploading to SharePoint...",
        req.filename, docx_size,
    )

    # Step 2: Upload raw binary to SharePoint via MS Graph
    upload_url = (
        f"https://graph.microsoft.com/v1.0/drives/{req.driveId}"
        f"/items/{req.folderId}:/{req.filename}:/content"
    )

    graph_req = urllib.request.Request(
        upload_url,
        data=docx_bytes,
        method="PUT",
        headers={
            "Authorization": f"Bearer {req.accessToken}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Length": str(docx_size),
        },
    )

    try:
        with urllib.request.urlopen(graph_req, timeout=30) as resp:
            resp_body = resp.read().decode("utf-8")
            upload_result = json.loads(resp_body)
            logger.info("SharePoint upload success: %s", upload_result.get("webUrl", ""))
    except urllib.error.HTTPError as e:
        error_body = e.read().decode("utf-8", errors="replace")
        logger.error("SharePoint upload failed: %s %s", e.code, error_body)
        raise HTTPException(
            502,
            f"SharePoint upload failed ({e.code}): {error_body[:500]}",
        )
    except Exception as e:
        logger.error("SharePoint upload error: %s", e)
        raise HTTPException(502, f"SharePoint upload error: {e}")

    return {
        "status": "uploaded",
        "filename": req.filename,
        "sizeBytes": docx_size,
        "sharePointUrl": upload_result.get("webUrl", ""),
        "sharePointItemId": upload_result.get("id", ""),
    }
